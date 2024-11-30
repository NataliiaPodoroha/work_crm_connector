from docx import Document
from io import BytesIO


GENERATED_CV_FILE_NAME = "cv_from_text.docx"


def create_cv_document(text):
    doc = Document()
    doc.add_heading("CV", level=1)

    paragraphs = text.replace('\r\n', '\n').split('\n')

    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream, GENERATED_CV_FILE_NAME
